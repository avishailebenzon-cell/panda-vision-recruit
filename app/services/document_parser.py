import logging
import io
from typing import Tuple, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service for parsing documents (PDF, DOCX, DOC)."""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".doc"}

    @staticmethod
    def parse_pdf(content: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            import fitz

            pdf_stream = io.BytesIO(content)
            document = fitz.open(stream=pdf_stream, filetype="pdf")

            text = ""
            for page_num, page in enumerate(document):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.get_text()

            document.close()
            logger.info(f"Successfully parsed PDF ({len(text)} characters)")
            return text

        except ImportError:
            logger.error("PyMuPDF (fitz) not installed. Install with: pip install PyMuPDF")
            raise
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise

    @staticmethod
    def parse_docx(content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document

            doc_stream = io.BytesIO(content)
            doc = Document(doc_stream)

            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            logger.info(f"Successfully parsed DOCX ({len(text)} characters)")
            return text

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise

    @staticmethod
    def parse_doc(content: bytes) -> str:
        """
        Extract text from old DOC format.
        Uses textract or fallback method.
        """
        try:
            # Try using textract first
            try:
                import textract
                text = textract.process(io.BytesIO(content), method="docx")
                if isinstance(text, bytes):
                    text = text.decode("utf-8", errors="ignore")
                logger.info(f"Successfully parsed DOC using textract ({len(text)} characters)")
                return text
            except ImportError:
                logger.warning("textract not installed, trying alternative method")
                # Fallback: try to parse as DOCX (some .doc files are actually DOCX format)
                return DocumentParser.parse_docx(content)

        except Exception as e:
            logger.error(f"Error parsing DOC: {e}")
            raise

    @staticmethod
    def parse_document(content: bytes, filename: str) -> Tuple[str, bool]:
        """
        Parse document content based on file extension.

        Args:
            content: Binary content of the document
            filename: Filename with extension

        Returns:
            Tuple of (extracted_text, success)
        """
        try:
            file_ext = Path(filename).suffix.lower()

            if file_ext == ".pdf":
                text = DocumentParser.parse_pdf(content)
            elif file_ext == ".docx":
                text = DocumentParser.parse_docx(content)
            elif file_ext == ".doc":
                text = DocumentParser.parse_doc(content)
            else:
                logger.warning(f"Unsupported file format: {file_ext}")
                return "", False

            # Clean up whitespace
            text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
            return text, True

        except Exception as e:
            logger.error(f"Failed to parse document {filename}: {e}")
            return "", False

    @staticmethod
    def is_supported_format(filename: str) -> bool:
        """Check if file format is supported."""
        file_ext = Path(filename).suffix.lower()
        return file_ext in DocumentParser.SUPPORTED_FORMATS
